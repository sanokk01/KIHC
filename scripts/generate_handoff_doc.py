from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GOLD = "9B7A37"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "D4DCE5"
WHITE = "FFFFFF"
BLACK = "1F2937"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size=11, color=BLACK, bold=False, italic=False, name="Calibri"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), "Malgun Gothic")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size=5):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def table_widths(column_count):
    if column_count == 2:
        return [2300, 7060]
    if column_count == 3:
        return [1700, 5200, 2460]
    if column_count == 4:
        return [1500, 3700, 1800, 2360]
    base = CONTENT_DXA // column_count
    values = [base] * column_count
    values[-1] += CONTENT_DXA - sum(values)
    return values


def add_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    widths = table_widths(len(rows[0]))
    for row_idx, values in enumerate(rows):
        for col_idx, value in enumerate(values):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value.strip())
            set_run_font(run, size=9.2 if len(rows[0]) >= 4 else 9.8, bold=row_idx == 0)
            if row_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
    mark_repeat_header(table.rows[0])
    set_table_geometry(table, widths)
    set_table_borders(table)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(3)


def add_numbering_definition(doc, ordered):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


def add_inline_runs(paragraph, text):
    for part in re.split(r"(`[^`]+`)", text):
        if not part:
            continue
        is_code = part.startswith("`") and part.endswith("`")
        value = part[1:-1] if is_code else part
        run = paragraph.add_run(value)
        set_run_font(
            run,
            name="Consolas" if is_code else "Calibri",
            size=9.5 if is_code else 11,
            color=DARK_BLUE if is_code else BLACK,
        )


def add_list_item(doc, text, num_id):
    p = doc.add_paragraph()
    apply_numbering(p, num_id)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    add_inline_runs(p, text)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.keep_together = False
    add_inline_runs(p, text)
    return p


def add_callout(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.2
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=DARK_BLUE)


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.right_indent = Inches(0.16)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.1
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F4F7")
    p_pr.append(shd)
    run = p.add_run("\n".join(lines))
    set_run_font(run, name="Consolas", size=8.7, color="344054")


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    hp = section.header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("KIHC  |  웹사이트 개선·운영 가이드")
    set_run_font(hr, size=8.5, color=MUTED, bold=True)

    fp = section.footer.paragraphs[0]
    fp.text = ""
    add_page_number(fp)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("PROJECT HANDOFF & OPERATIONS GUIDE")
    set_run_font(r, size=9.5, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("KIHC 웹사이트 개선 작업 및 운영 적용 가이드")
    set_run_font(r, size=28, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("모바일 최적화 · 관리자 보안 · Supabase/Netlify 운영 절차")
    set_run_font(r, size=13.5, color=MUTED)

    metadata = [
        ("문서 버전", "1.0"),
        ("작성일", "2026-08-12"),
        ("전달 대상", "기획팀 · 콘텐츠 담당 · 운영/배포 담당"),
        ("보안 등급", "내부 공유용 · 비밀번호 원문 미포함"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(f"{label}  ")
        set_run_font(r, size=9, color=GOLD, bold=True)
        r = p.add_run(value)
        set_run_font(r, size=10, color=NAVY, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_markdown(doc, markdown):
    lines = markdown.splitlines()
    index = 0
    while index < len(lines) and not lines[index].startswith("> "):
        index += 1
    if index < len(lines):
        add_callout(doc, lines[index][2:].strip())
        index += 1

    in_code = False
    code_lines = []
    list_kind = None
    list_num_id = None

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        if stripped == "---":
            list_kind = None
            list_num_id = None
            index += 1
            continue

        if stripped.startswith("|") and index + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+", lines[index + 1]):
            rows = []
            header = [x.strip() for x in stripped.strip("|").split("|")]
            rows.append(header)
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append([x.strip().replace("`", "") for x in lines[index].strip().strip("|").split("|")])
                index += 1
            add_table(doc, rows)
            list_kind = None
            list_num_id = None
            continue

        heading = re.match(r"^(#{2,3})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1)) - 1
            p = doc.add_paragraph(heading.group(2), style=f"Heading {level}")
            p.paragraph_format.keep_with_next = True
            list_kind = None
            list_num_id = None
            index += 1
            continue

        bullet = re.match(r"^-\s+(.+)$", stripped)
        ordered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet or ordered:
            kind = "ordered" if ordered else "bullet"
            if list_kind != kind:
                list_num_id = add_numbering_definition(doc, ordered=kind == "ordered")
                list_kind = kind
            add_list_item(doc, (ordered or bullet).group(1), list_num_id)
            index += 1
            continue

        if stripped.startswith("> "):
            add_callout(doc, stripped[2:])
            index += 1
            continue

        add_body_paragraph(doc, stripped)
        list_kind = None
        list_num_id = None
        index += 1


def build(markdown_path, output_path):
    markdown = Path(markdown_path).read_text(encoding="utf-8")
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_cover(doc)
    parse_markdown(doc, markdown)
    doc.core_properties.title = "KIHC 웹사이트 개선 작업 및 운영 적용 가이드"
    doc.core_properties.subject = "기획팀·운영팀 전달용 작업 및 배포 가이드"
    doc.core_properties.author = "KIHC Web Project"
    doc.core_properties.keywords = "KIHC, Supabase, Netlify, 모바일 최적화, 관리자"
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: generate_handoff_doc.py INPUT.md OUTPUT.docx")
    build(sys.argv[1], sys.argv[2])
